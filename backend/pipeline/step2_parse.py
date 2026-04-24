import io
import re
from typing import Optional

import httpx

from config import settings
from logger import get_logger

logger = get_logger(__name__)


def _normalize_text(text: str) -> str:
    text = re.sub(r"\r", "\n", text or "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


async def _download_pdf(url: str) -> bytes:
    if url.startswith("mock://"):
        raise ValueError("Raw PDF bytes are required for mock Cloudinary URLs")
    async with httpx.AsyncClient(timeout=settings.api_timeout_seconds) as client:
        response = await client.get(url)
        response.raise_for_status()
        return response.content


def _extract_with_pymupdf(raw_pdf_bytes: bytes) -> tuple[str, int]:
    import fitz

    doc = fitz.open(stream=raw_pdf_bytes, filetype="pdf")
    try:
        text = "\n".join(page.get_text("text") for page in doc)
        return text, doc.page_count
    finally:
        doc.close()


def _extract_with_pdfplumber(raw_pdf_bytes: bytes) -> tuple[str, int]:
    import pdfplumber

    with pdfplumber.open(io.BytesIO(raw_pdf_bytes)) as pdf:
        text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        return text, len(pdf.pages)


def _extract_with_docx(raw_bytes: bytes) -> tuple[str, int]:
    from docx import Document

    document = Document(io.BytesIO(raw_bytes))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    table_cells = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    table_cells.append(cell.text)
    return "\n".join(paragraphs + table_cells), 1


async def extract_text_from_resume(
    *,
    resume_id: str,
    filename: str = "resume.pdf",
    raw_pdf_bytes: Optional[bytes] = None,
    cloudinary_url: Optional[str] = None,
) -> str:
    logger.info("[STEP 2] Parsing resume with Python libraries only...")
    if raw_pdf_bytes is None:
        if not cloudinary_url:
            raise ValueError("Either raw_pdf_bytes or cloudinary_url is required")
        raw_pdf_bytes = await _download_pdf(cloudinary_url)

    text = ""
    pages = 0
    lower_filename = filename.lower()
    if lower_filename.endswith(".docx"):
        logger.info("[INFO] Using python-docx parser | resume_id=%s", resume_id)
        text, pages = _extract_with_docx(raw_pdf_bytes)
        normalized_docx = _normalize_text(text)
        if not normalized_docx:
            raise ValueError("No text could be extracted from resume DOCX")
        logger.info("[SUCCESS] Resume parsed | parser=python-docx | chars=%s", len(normalized_docx))
        return normalized_docx

    if lower_filename.endswith(".txt"):
        logger.info("[INFO] Using plain text parser | resume_id=%s", resume_id)
        normalized_txt = _normalize_text(raw_pdf_bytes.decode("utf-8", errors="ignore"))
        if not normalized_txt:
            raise ValueError("No text could be extracted from resume TXT")
        logger.info("[SUCCESS] Resume parsed | parser=text | chars=%s", len(normalized_txt))
        return normalized_txt

    try:
        logger.info("[INFO] Using PyMuPDF parser | resume_id=%s", resume_id)
        text, pages = _extract_with_pymupdf(raw_pdf_bytes)
        if len(text.strip()) < 100:
            logger.warning("[WARN] PyMuPDF returned <100 chars — switching to pdfplumber")
            raise ValueError("PyMuPDF extracted too little text")
        logger.info("[SUCCESS] Extracted %s pages | %s characters", pages, len(text))
    except Exception as first_error:
        try:
            logger.info("[INFO] pdfplumber fallback active | resume_id=%s", resume_id)
            text, pages = _extract_with_pdfplumber(raw_pdf_bytes)
            logger.info("[SUCCESS] pdfplumber extracted %s characters", len(text))
        except Exception as second_error:
            logger.warning(
                "[WARN] PDF parsers failed — attempting text decode fallback | resume_id=%s | error=%s",
                resume_id,
                second_error or first_error,
            )
            text = raw_pdf_bytes.decode("utf-8", errors="ignore")
            pages = 1

    normalized = _normalize_text(text)
    logger.debug("[DEBUG] Whitespace normalized | final_length=%s", len(normalized))
    if not normalized:
        raise ValueError("No text could be extracted from resume PDF")
    logger.info("[SUCCESS] Extracted %s pages | %s characters", pages, len(normalized))
    return normalized


async def extract_text_from_pdf(
    *,
    resume_id: str,
    raw_pdf_bytes: Optional[bytes] = None,
    cloudinary_url: Optional[str] = None,
) -> str:
    return await extract_text_from_resume(
        resume_id=resume_id,
        filename="resume.pdf",
        raw_pdf_bytes=raw_pdf_bytes,
        cloudinary_url=cloudinary_url,
    )
